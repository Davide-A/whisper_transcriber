import os
import time
import whisper
import tiktoken
import torch
from docx import Document
from pathlib import Path
from pydub import AudioSegment
from transformers import AutoTokenizer, AutoModelForCausalLM


def chunk_audio(input_file, temp_dir, chunk_size_mb=20):
    chunk_size_bytes = chunk_size_mb * 1024 * 1024
    audio = AudioSegment.from_file(input_file)

    audio_size = os.path.getsize(input_file)
    chunk_duration_ms = len(audio) * (chunk_size_bytes / audio_size)

    os.makedirs(temp_dir, exist_ok=True)

    for i in range(0, len(audio), int(chunk_duration_ms)):
        chunk = audio[i:i + int(chunk_duration_ms)]
        chunk_file = os.path.join(temp_dir, f"chunk_{i // int(chunk_duration_ms)}.mp3")
        chunk.export(chunk_file, format="mp3")
        yield chunk_file


def milliseconds_until_sound(sound, silence_threshold_in_decibels=-20.0, chunk_size=10):
    trim_ms = 0  # ms

    assert chunk_size > 0  # to avoid infinite loop
    while sound[trim_ms:trim_ms+chunk_size].dBFS < silence_threshold_in_decibels and trim_ms < len(sound):
        trim_ms += chunk_size

    return trim_ms


def trim_start(filepath, dir):
    path = Path(filepath)
    directory = Path(dir)
    filename = path.name

    if filename.endswith(".WAV"):
        audio = AudioSegment.from_file(filepath, format="WAV")
        start_trim = milliseconds_until_sound(audio)
        trimmed = audio[start_trim:]
        new_filename = directory / f"trimmed_{filename}"
        trimmed.export(new_filename, format="WAV")

    elif filename.endswith(".mp3"):
        audio = AudioSegment.from_file(filepath, format="mp3")
        start_trim = milliseconds_until_sound(audio)
        trimmed = audio[start_trim:]
        new_filename = directory / f"trimmed_{filename}"
        trimmed.export(new_filename, format="mp3")

    return trimmed, new_filename


def remove_non_ascii(text):
    return ''.join(i for i in text if ord(i)<128)


def transcribe_audio(temp_dir, file):

    print(f"processing file:{file}")

    prompt = (
        f"Questo è il frammento di una lezione universitaria di teologia. L'argomento "
        f"è trattato attraerso le fonti storiche. Fai attenzione ai nomi biblici dell'antico testamento come:"
        f"Samuele, Saul, Rut, Davide, Salomone, Roboamo, Geroboamo."
    )

    audio_path = os.path.join(temp_dir, file)
    
    transcription = model.transcribe(audio_path,
                                     language='it',
                                     verbose=True,
                                     initial_prompt=prompt)
    
    return transcription['text']


def biblical_assistant(input_text):

    system_prompt = f"""Sei un dotto assistente universitario esperto di teologia e antico testamente.
    Ti sarà fornita la trascrizione di una lezione universitaria di teologia. Il tuo compito è revisionarla, correggendo la sintassi della transcrizione, aggiungere la punteggiatura necessaria per produrre un discorso scorrevole. 
    Dovresti aggiungere quindi, virgole, punti alla fine dei periodi, punti di domanda, lettere maiusole e formattare il discorso correttamente. Considera che la trascrizione potrebbe contenere ripetizioni e allucinazioni del modello che l'ha trascritta.
    Inoltre dovrai correggere eventuali errori nei nomi storici, biblici dell'antico testamento che dovessero essere stati trascritti male, completare eventuali discorsi incompleti a causa di errori di trascrizione, approfondendo l'argomento con quello che sai per produrre un testo scorrevole e comprensibile sul quale 
    poi permettere agli studenti di studiare. Cerca non non tralasciare nulla, i dettagli e aneddoti storici possono essere importanti per eccellere nell'esame.
    La trascrizione ti sarà fornita in più parti, crea un titolo attinente all'argomento per ogni parte.
    Adesso formatta questo testo:
    {input_text}
    """

    inputs = tokenizer(system_prompt, return_tensors="pt", truncation=True, max_length=2048)

    outputs = model.generate(**inputs, max_length=4096, temperature=0.9)
    
    formatted_text = tokenizer.decode(outputs[0], skip_special_tokens=True)

    return formatted_text


def calculate_tokens(text, model="gpt-4"):
    tokenizer = tiktoken.encoding_for_model(model)
    return len(tokenizer.encode(text))


def split_text(text, max_tokens, model="gpt-4"):
    tokenizer = tiktoken.encoding_for_model(model)
    tokens = tokenizer.encode(text)
    for i in range(0, len(tokens), max_tokens):
        yield tokenizer.decode(tokens[i:i+max_tokens])


def split_text_with_overlap(text, chunk_size, overlap_size):
    """
    Split the text into chunks with overlapping sections.

    Parameters:
    - text: The input text to split.
    - chunk_size: The maximum size of each chunk.
    - overlap_size: The number of characters to overlap between chunks.

    Returns:
    - List of chunks with overlaps.
    """
    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap_size  # Move forward with overlap
        
    return chunks
    

# Modify the transcription loop
def process_transcription(ascii_transcript, output_dir_assis, model, chunk_size=6500, overlap_size=200):

    assistant_model
    print("Invoking biblical assistant...")
    chunks = split_text_with_overlap(ascii_transcript, chunk_size, overlap_size)

    with open(output_dir_assis, "w", encoding="utf-8") as f:
        for idx, chunk in enumerate(chunks):
            print(f"Processing chunk {idx + 1}/{len(chunks)}...")
            # Send chunk to the model
            final_transcript = biblical_assistant(chunk)

            # Write the final transcript to the file
            f.write(final_transcript + "\n")  # Add spacing between chunks
            print(f"Chunk {idx + 1} written to file.")


if __name__ == '__main__':

    # intro sacra scrittura
    # teologia morale fondamentale
    # teologia fondamentale
    # storia della chiesa 1-4

    # corso = "teo_mor_fond"
    # corso = "int_scrit"
    # corso = "st_chiesa"
    corso = "teo_fond"
    
    input_dir = f"media/input/{corso}/"
    output_dir_trans = f"media/output/{corso}/transcripted"
    output_dir_assis = f"media/output/{corso}/processed"
    temp_trim = "media/temp_trimmed"
    temp_dir = "media/temp"

    # Initialize Whisper model and assistant
    print("Inizialising the models")
    tokenizer = AutoTokenizer.from_pretrained("tiiuae/falcon-11B")
    assistant_model = AutoModelForCausalLM.from_pretrained("tiiuae/falcon-11B", device_map="auto")
    tokenizer = AutoTokenizer.from_pretrained(assistant_model)
    model = whisper.load_model("large")
    print("\n")

    # Trim the start of the original audio file to remove silence from audio
    print("Trimming the audio file")
    audios = os.listdir(input_dir)
    trimmed_audio, trimmed_filename = trim_start(input_dir + audios[0], temp_trim)
    if trimmed_filename.suffix == ".WAV":
        trimmed_audio = AudioSegment.from_file(trimmed_filename, format="WAV")
        format = "WAV"
    elif trimmed_filename.suffix == ".mp3":
        trimmed_audio = AudioSegment.from_file(trimmed_filename, format="mp3")
        format = "mp3"
    print("\n")

    # here the aduio file is divided in chunks
    one_minute = 1 * 60 * 1000  # Duration for each segment (in milliseconds)

    start_time = 0  # Start time for the first segment

    i = 0  # Index for naming the segmented files

    if not os.path.isdir(temp_dir):  # Create the output directory if it does not exist
        os.makedirs(temp_dir)

    while start_time < len(trimmed_audio):  # Loop over the trimmed audio file
        segment = trimmed_audio[start_time:start_time + one_minute]  # Extract a segment
        segment.export(os.path.join(temp_dir, f"trimmed_{i:02d}.{format}"), format=format)  # Save the segment
        start_time += one_minute  # Update the start time for the next segment
        i += 1  # Increment the index for naming the next file

    # Get list of trimmed and segmented audio files and sort them numerically
    audio_files = sorted(
        (f for f in os.listdir(temp_dir)),
        key=lambda f: int(''.join(filter(str.isdigit, f)))
    )

    print("Starting transcription")
    transcription_start = time.time()
    transcriptions = [transcribe_audio(temp_dir, file) for file in audio_files]
    full_transcript = ' '.join(transcriptions)
    print(f"Transciption made in: {time.time()-transcription_start} seconds")

    if not os.path.isdir(output_dir_trans):  # Create the output directory if it does not exist
        os.makedirs(output_dir_trans)

    file_name = f"{Path(audios[0]).stem}_transcripted.txt"
    output_file_path  = os.path.join(output_dir_trans, file_name)

    with open(output_file_path, 'w', encoding='utf-8') as trans:
        trans.write(full_transcript)
    print("\n")
    print("Transcription saved")

    # cleaning temp
    files = os.listdir(temp_dir)
    trimmed = os.listdir(temp_trim)
    temp = files + trimmed
    for file in temp:
        file_path = os.path.join(temp_dir, file)
        if os.path.isfile(file_path):
            os.remove(file_path)
    print("\n")
    print(full_transcript[0:1000])

    # Remove non-ascii characters from the transcript
    print("Removing non ascii characters...")
    print(f"Tokens in input: {calculate_tokens(full_transcript)}")
    ascii_transcript = remove_non_ascii(full_transcript)
    print("\n")

    if not os.path.isdir(output_dir_assis):  # Create the output directory if it does not exist
        os.makedirs(output_dir_assis)

    file_name = f"{Path(audios[0]).stem}_processed.txt"
    output_file_path  = os.path.join(output_dir_assis, file_name)
    process_transcription(ascii_transcript, output_file_path, pipeline)

    with open(output_file_path, 'r', encoding='utf-8') as file:
        text_content_new = file.read()

    doc_new = Document()    
    path = Path(output_file_path)
    filename = path.name

    # suddivisone del testo in paragrafi
    paragraphs_new = text_content_new.split("\n\n")

    for paragraph in paragraphs_new:

        if paragraph.startswith("Titolo:"):
            title = paragraph.replace("Titolo:", "").strip()
            doc_new.add_heading(title, level=2)
        else:
            doc_new.add_paragraph(paragraph.strip())

    # Save the document to a Word file
    output_path_new = f"{filename}"
    doc_new.save(output_path_new)
    print("Transcription saved. Ending job")
