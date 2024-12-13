import os
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document

def generate_lecture_notes(content):

    with open("prompts/notes.txt", "r") as file:
        prompt_template = file.read()

    prompt = prompt_template.format(content=content)

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Sei un assistente che aiuta studenti a generare appunti approfonditi e dettagliat dalla transcrizione di audio lezioni."},
                {"role": "user", "content": prompt},
            ],
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"An error occurred: {e}")
        return None


if __name__ == "__main__":

    text_dir = "media/output/2_Intro_Sacra_scrittura_Grassilli_transcription.txt"
    notes = "media/notes/lecture_notes.txt"

    load_dotenv()
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    with open(text_dir, 'r', encoding="utf-8") as text:
        transcription = text.read()

    lecture_notes = generate_lecture_notes(transcription)

    #doc = Document()
    #doc.add_heading("Generated Text", level=1)
    #doc.add_paragraph(lecture_notes)
    #doc.save(notes)

    if lecture_notes:
        print("Lecture Notes:")
        print(lecture_notes)
    
        with open(notes, "w") as file:
            file.write(lecture_notes)
    
        print(f"Lecture notes saved to {notes}")
    else:
        print("Failed to generate lecture notes.")